import warnings
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import os
from docx.enum.table import WD_ALIGN_VERTICAL

def apply_styles(element, run):
    # Apply styles based on HTML elements
    if element.name == 'b':
        run.bold = True
    elif element.name == 'i':
        run.italic = True
    elif element.name == 'u':
        run.underline = True
    elif element.name == 'font':
        if 'color' in element.attrs:
            run.font.color.rgb = hex_to_rgb(element['color'])
        if 'size' in element.attrs:
            run.font.size = Pt(float(element['size']))
        if 'face' in element.attrs:
            run.font.name = element['face']
    elif element.name == 'a':
        run.underline = True
        if 'href' in element.attrs:
            run.hyperlink.address = element['href']
            run.hyperlink.target_mode = True
    # Add more styles as needed

def handle_flex_css(element, paragraph):
    # Handle flex CSS properties
    if 'flex' in element.attrs:
        flex_value = element['flex']
        if flex_value == '1':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif flex_value == '2':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def handle_grid_css(element, table):
    # Handle grid CSS properties by creating a table
    if 'grid-template-columns' in element.attrs:
        columns = element['grid-template-columns'].split()
        for col in columns:
            table.add_column(Inches(float(col)))

def handle_images(element, parent):
    # Handle image elements
    if element.name == 'img' and 'src' in element.attrs:
        image_src = element['src']
        image_path = download_image(image_src)
        parent.add_picture(image_path, width=Inches(3))  # Adjust the width as needed
        os.remove(image_path)

def download_image(image_url):
    # You'll need to implement image downloading logic based on your specific use case
    # This example assumes the image is downloaded and saved to a temp file
    # Replace this with your actual image download logic
    # For simplicity, we'll just save a placeholder image
    placeholder_image_path = 'placeholder.jpg'  # Replace with your image path or URL
    temp_image_path = 'temp_image.jpg'
    with open(temp_image_path, 'wb') as img_file:
        with open(placeholder_image_path, 'rb') as placeholder_file:
            img_file.write(placeholder_file.read())
    return temp_image_path

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def convert_html_to_docx(html_file, output_path):
    """Converts an HTML file to a DOCX file.

    Args:
        html_file (str): The path to the HTML file.
        output_path (str): The path for the output DOCX file.
    """

    # Parse the HTML file.
    soup = BeautifulSoup(open(html_file, "r", encoding="utf-8").read(), "html.parser")

    # Create a new DOCX document.
    document = Document()

    def handle_horizontal_line(document):
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)  # Add a page break to simulate a horizontal line

    # Recursively add elements to the DOCX document.
    def add_element_to_docx(element, parent):
        if element.name == 'p':
            paragraph = parent.add_paragraph()
            if paragraph.runs:
                apply_styles(element, paragraph.runs[0])
            handle_flex_css(element, paragraph)
            handle_images(element, paragraph)
            for child in element.children:
                run = paragraph.add_run(child.text)
                apply_styles(child, run)
        elif element.name in ['h1', 'h2', 'h3']:
            heading = parent.add_heading(element.text, level=int(element.name[1]))
        elif element.name in ['ul', 'ol']:
            list_style = 'ListBullet' if element.name == 'ul' else 'ListNumber'
            list_paragraph = parent.add_paragraph(style=list_style)
            for child in element.find_all("li"):
                add_element_to_docx(child, list_paragraph)
        elif element.name == 'table':
            rows = int(element.get("rows", 1))
            cols = int(element.get("cols", 1))
            table = parent.add_table(rows, cols)
            table.alignment = WD_ALIGN_VERTICAL.CENTER
            handle_grid_css(element, table)
            for row in element.find_all("tr"):
                for cell in row.find_all(["th", "td"]):
                    add_element_to_docx(cell, table.cell(0, 0))
        elif element.name == 'hr':
            handle_horizontal_line(document)

    # Add the HTML content to the DOCX document.
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'table', 'hr']):
        add_element_to_docx(element, document)

    # Save the DOCX document.
    document.save(output_path)

# Example usage
html_file_path = "E:/c++/harsh/index.html"  # Replace with your HTML file path
docx_file_path = "output.docx"   # Replace with your desired output DOCX file path

# Suppress the specific UserWarning
warnings.filterwarnings("ignore", category=UserWarning, module="docx")

convert_html_to_docx(html_file_path, docx_file_path)

# Restore the warning behavior if needed
# warnings.filterwarnings("default", category=UserWarning, module="docx")



