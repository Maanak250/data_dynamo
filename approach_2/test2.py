from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

html_file_path = "E:/c++/harsh/index.html"
output_file_path = "output.docx"

# Read the HTML content from the file
with open(html_file_path, "r", encoding="utf-8") as html_file:
    html_content = html_file.read()

# Parse the HTML content
soup = BeautifulSoup(html_content, 'html.parser')

def apply_css_styles(css_styles, element):
    if 'font-size' in css_styles:
        element.runs[0].font.size = Pt(float(css_styles['font-size']))
    if 'color' in css_styles:
        element.runs[0].font.color.rgb = RGBColor.from_string(css_styles['color'])
    if 'font-family' in css_styles:
        element.runs[0].font.name = css_styles['font-family']

def process_paragraph(paragraph, element):
    apply_css_styles(element.get('style', {}), paragraph)
    paragraph.add_run(element.get_text())

def process_heading(heading, element):
    level = int(element.name[1])
    apply_css_styles(element.get('style', {}), heading)
    heading.add_run(element.get_text())

def process_unordered_list(ul, element):
    ul = document.add_paragraph()
    ul.style = 'List Bullet'
    for li in element.find_all('li'):
        apply_css_styles(li.get('style', {}), ul)
        ul.add_run('• ' + li.get_text())

def process_ordered_list(ol, element):
    ol = document.add_paragraph()
    ol.style = 'List Number'
    for li in element.find_all('li'):
        apply_css_styles(li.get('style', {}), ol)
        ol.add_run(str(li.find_previous_siblings('li').count(li) + 1) + '. ' + li.get_text())

def process_line_break(document, element):
    document.add_paragraph('')

def process_horizontal_rule(document, element):
    document.add_paragraph('____________________________________________________________________________')

def process_italic(italic, element):
    i = document.add_paragraph()
    apply_css_styles(element.get('style', {}), i)
    i.add_run(element.get_text()).italic = True

def process_bold(bold, element):
    b = document.add_paragraph()
    apply_css_styles(element.get('style', {}), b)
    b.add_run(element.get_text()).bold = True

# Create a new DOCX document
doc = Document()

for element in soup.recursiveChildGenerator():
    if element.name == 'p':
        process_paragraph(doc.add_paragraph(), element)
    elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        process_heading(doc.add_heading(), element)
    elif element.name == 'ul':
        process_unordered_list(doc, element)
    elif element.name == 'ol':
        process_ordered_list(doc, element)
    elif element.name == 'br':
        process_line_break(doc, element)
    elif element.name == 'hr':
        process_horizontal_rule(doc, element)
    elif element.name == 'i':
        process_italic(doc.add_paragraph(), element)
    elif element.name == 'b':
        process_bold(doc.add_paragraph(), element)

doc.save(output_file_path)
