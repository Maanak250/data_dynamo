from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

html_file_path = "E:/c++/harsh/index.html"
output_file_path = "output.docx"

# Read the HTML content from the file
with open(html_file_path, "r", encoding="utf-8") as html_file:
    html_content = html_file.read()

# Parse the HTML content
soup = BeautifulSoup(html_content, 'html.parser')

def apply_css_styles(css_styles, element):
    if 'font-size' in css_styles:
        element.runs[0].font.size = Pt(float(css_styles['font-size']))
    if 'color' in css_styles:
        element.runs[0].font.color.rgb = RGBColor.from_string(css_styles['color'])
    if 'font-family' in css_styles:
        element.runs[0].font.name = css_styles['font-family']

def convert_html_to_docx(soup, document):
    for element in soup.recursiveChildGenerator():
        if element.name == 'p':
            p = document.add_paragraph()
            apply_css_styles(element.get('style', {}), p)
            p.add_run(element.get_text())
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = document.add_heading(level=level)
            apply_css_styles(element.get('style', {}), heading)
            heading.add_run(element.get_text())
        elif element.name == 'ul':
            ul = document.add_paragraph()
            ul.style = 'List Bullet'
            for li in element.find_all('li'):
                apply_css_styles(li.get('style', {}), ul)
                ul.add_run('• ' + li.get_text())
        elif element.name == 'ol':
            ol = document.add_paragraph()
            ol.style = 'List Number'
            for li in element.find_all('li'):
                apply_css_styles(li.get('style', {}), ol)
                ol.add_run(str(li.find_previous_siblings('li').count(li) + 1) + '. ' + li.get_text())
        elif element.name == 'br':
            document.add_paragraph('')
        elif element.name == 'hr':
            document.add_paragraph().add_run('____________________________________________________________________________').bold = True
        elif element.name == 'i':
            i = document.add_paragraph()
            apply_css_styles(element.get('style', {}), i)
            i.add_run(element.get_text()).italic = True
        elif element.name == 'b':
            apply_css_styles(element.get('style', {}), p)
            p.add_run(element.get_text()).bold = True

# Create a new DOCX document
doc = Document()
convert_html_to_docx(soup, doc)
doc.save(output_file_path)


