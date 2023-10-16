from bs4 import BeautifulSoup
import cssutils
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

# from shared import Pt, RGBColor

class HTMLComponent:
    def _init_(self, css_styles):
        self.css_styles = css_styles
        self.font_size = self.css_styles.get('font-size', None)
        self.color = self.css_styles.get('color', None)
        self.font_family = self.css_styles.get('font-family', None)

    def apply_custom_styles(self):
        pass

class Paragraph(HTMLComponent):
    def init(self, css_styles, text, font_size, color, font_family):
        super().init(css_styles)
        self.text = text
        self.font_size = font_size
        self.color = color
        self.font_family = font_family

    def apply_custom_styles(self, paragraph):
        if self.font_size:
            paragraph.runs[0].font.size = Pt(float(self.font_size))
        if self.color:
            paragraph.runs[0].font.color.rgb = RGBColor(*tuple(int(self.color[i:i+2], 16) for i in (1, 3, 5)))
        if self.font_family:
            paragraph.runs[0].font.name = self.font_family

class Header(HTMLComponent):
    def init(self, css_styles, tag_name, text, font_size, color, font_family):
        super().init(css_styles)
        self.tag_name = tag_name
        self.text = text
        self.font_size = font_size
        self.color = color
        self.font_family = font_family

    def apply_custom_styles(self, paragraph):
        if self.font_size:
            paragraph.runs[0].font.size = Pt(float(self.font_size))
        if self.color:
            paragraph.runs[0].font.color.rgb = RGBColor(*tuple(int(self.color[i:i+2], 16) for i in (1, 3, 5)))
        if self.font_family:
            paragraph.runs[0].font.name = self.font_family

class Footer(HTMLComponent):
    def init(self, css_styles, text, color, font_family):
        super().init(css_styles)
        self.text = text
        self.color = color
        self.font_family = font_family

    def apply_custom_styles(self, paragraph):
        if self.color:
            paragraph.runs[0].font.color.rgb = RGBColor(*tuple(int(self.color[i:i+2], 16) for i in (1, 3, 5)))
        if self.font_family:
            paragraph.runs[0].font.name = self.font_family

class Table(HTMLComponent):
    def init(self, css_styles, html, font_size, color, font_family):
        super().init(css_styles)
        self.html = html
        self.font_size = font_size
        self.color = color
        self.font_family = font_family

    def apply_custom_styles(self, table):
        if self.font_size:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].font.size = Pt(float(self.font_size))
        if self.color:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].font.color.rgb = RGBColor(*tuple(int(self.color[i:i+2], 16) for i in (1, 3, 5)))
        if self.font_family:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].font.name = self.font_family

class Heading(HTMLComponent):
    def init(self, css_styles, text, font_size, color, font_family):
        super().init(css_styles)
        self.text = text
        self.font_size = font_size
        self.color = color
        self.font_family = font_family

    def apply_custom_styles(self, paragraph):
        if self.font_size:
            paragraph.runs[0].font.size = Pt(float(self.font_size))
        if self.color:
            paragraph.runs[0].font.color.rgb = RGBColor(*tuple(int(self.color[i:i+2], 16) for i in (1, 3, 5)))
        if self.font_family:
            paragraph.runs[0].font.name = self.font_family

def extract_html_and_css(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        html_string = file.read()

    soup = BeautifulSoup(html_string, 'html.parser')

    extracted_html = soup.prettify()

    style_tag = soup.head.find('style')
    extracted_css = {}
    if style_tag:
        css_string = style_tag.get_text()
        css_rules = cssutils.parseString(css_string)
        for rule in css_rules:
            if isinstance(rule, cssutils.css.CSSComment):
                continue
            selector = rule.selectorText
            styles = {}
            for prop in rule.style:
                styles[prop.name] = prop.value
            extracted_css[selector] = styles

    return extracted_html, extracted_css

def create_docx(output_file_path, components):
    document = Document()
    for component in components:
        if isinstance(component, Paragraph):
            p = document.add_paragraph(style=None)
            p.add_run(component.text).font.size = Pt(component.font_size)
            p.add_run().font.color.rgb = RGBColor.from_string(component.color)
            p.style.font.name = component.font_family
        elif isinstance(component, Header):
            level = int(component.tag_name[-1])  # Convert level to integer
            header = document.add_heading(level=level)
            header.add_run(component.text).font.size = Pt(component.font_size)
            header.add_run().font.color.rgb = RGBColor.from_string(component.color)
            header.style.font.name = component.font_family
        elif isinstance(component, Table):
            table = document.add_table(rows=component.num_rows, cols=component.num_cols)
            for i, cell in enumerate(table._cells):
                cell.text = component.cells[i]
    document.save(output_file_path)

class HTMLComponentsExtractor:
    def _init_(self, html, css):
        self.html = html
        self.css = css
        self.soup = BeautifulSoup(self.html, 'html.parser')
        self.components = {
            'paragraphs': [],
            'headers': [],
            'footers': [],
            'tables': [],
            'headings': []
        }

    def get_element_styles(self, tag):
        css_styles = {}
        if tag.has_attr('class'):
            for class_ in tag['class']:
                css_styles.update(self.css.get(class_, {}))
        css_styles.update(self.css.get(f"{tag.name}.{tag.get('class')[0]}", {}))
        css_styles.update(self.css.get(tag.name, {}))
        return css_styles

    def extract_paragraphs(self):
        paragraphs = self.soup.find_all('p')
        for p in paragraphs:
            css_styles = self.get_element_styles(p)
            font_size = css_styles.get('font-size', None)
            color = css_styles.get('color', None)
            font_family = css_styles.get('font-family', None)
            text = p.get_text()
            paragraph = Paragraph(css_styles, text, font_size, color, font_family)
            self.components["paragraphs"].append(paragraph)

    def extract_headers(self):
        headers = self.soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        for header in headers:
            css_styles = self.get_element_styles(header)
            tag_name = header.name
            text = header.get_text()
            font_size = header.get('font-size', None)
            color = header.get('color', None)
            font_family = header.get('font-family', None)
            header = Header(css_styles, tag_name, text, font_size, color, font_family)
            self.components['headers'].append(header)

    def extract_footers(self):
        footers = self.soup.find_all(class_='footer')  # Use self.soup, not the_soup
        for footer in footers:
            # Extract and apply CSS styles (if available)
            css_style = footer.get('style', '') + self.get_css_style(footer)
            print(f"Footer: {footer.get_text()}\nCSS: {css_style}\n")


    def extract_tables(self):
        tables = self.soup.find_all('table')
        for table in tables:
            css_styles = self.get_element_styles(table)
            table_html = str(table)
            font_size = table.get('font-size', None)
            color = table.get('color', None)
            font_family = table.get('font-family', None)
            table = Table(css_styles, table_html, font_size, color, font_family)
            self.components['tables'].append(table)

    def extract_headings(self):
        headings = self.soup.find_all(class_='heading')
        for heading in headings:
            css_styles = self.get_element_styles(heading)
            text = heading.get_text()
            font_size = heading.get('font-size', None)
            color = heading.get('color', None)
            font_family = heading.get('font-family', None)
            heading = Heading(css_styles, text, font_size, color, font_family)
            self.components['headings'].append(heading)

    def parse_html(self):
        self.extract_paragraphs()
        self.extract_headers()
        self.extract_footers()
        self.extract_tables()
        self.extract_headings()

    def generate_docx_from_components(self, output_file_path):
        create_docx(output_file_path, self.components)

# Usage
file_path = "C:/Projects/convert/backend/uploads/sample_level_1.html"
output_file_path = "output.docx"
html_string, css_string = extract_html_and_css(file_path)

extractor = HTMLComponentsExtractor(html_string, css_string)
extractor.parse_html()
extractor.generate_docx_from_components(output_file_path)